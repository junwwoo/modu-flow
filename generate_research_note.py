"""
8주차 종합설계프로젝트 연구노트 생성 스크립트
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "맑은 고딕"


def set_run_font(run, name=FONT_NAME, size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_paragraph(doc, text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_code(doc, code_text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return p


def set_cell_text(cell, text, size=10, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_summary_table(doc):
    """1페이지: 요약 표"""
    add_heading(doc, "종합설계프로젝트 연구노트 (8주차)", size=16, bold=True)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(13)

    set_cell_text(table.cell(0, 0), "연구일시", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "2026년 04월 30일")

    set_cell_text(table.cell(1, 0), "연구장소", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 1), "경성대학교")

    set_cell_text(table.cell(2, 0), "수행자", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 1), "박준우, 이경민, 임용완, 정주영")

    set_cell_text(table.cell(3, 0), "연구내용", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    content_cell = table.cell(3, 1)
    content_cell.text = ""

    sections = [
        ("목적",
         "박준우 : 기존 웹캠 단일 스크립트(test_pose_7.py) 구조를 모바일 안드로이드 클라이언트에서 "
         "호출 가능한 모듈 + 서버 구조로 분리하고, REST 및 WebSocket 기반의 자세 분석 서버를 구축하였다."),
        ("연구방법",
         "박준우 : Python의 MediaPipe Tasks API와 OpenCV를 사용하여 단일 프레임 분석 함수를 모듈화하였고, "
         "FastAPI와 Uvicorn을 사용하여 REST(POST /analyze) 및 WebSocket(/ws) 엔드포인트를 구현하였다. "
         "base64 이미지 디코딩에는 Pillow를, CORS 처리에는 fastapi.middleware.cors를 활용하였다."),
        ("도출결과",
         "박준우 : analyze_pose(image) 함수를 핵심으로 하는 test_pose_8.py 모듈과, "
         "REST/WebSocket 두 가지 통신 방식을 모두 지원하는 pose_server.py FastAPI 서버를 완성하였다. "
         "WebSocket 핸들러는 asyncio.to_thread를 통해 CPU 바운드 추론을 분리함으로써 다중 클라이언트 동시 처리가 가능하도록 하였다."),
        ("문제점 분석",
         "박준우 : MediaPipe PoseLandmarker의 VIDEO 모드는 단조 증가 timestamp가 필요하여 "
         "단일 프레임 분석에 적합하지 않은 문제가 있었다. 또한 동기 함수인 analyze_pose를 비동기 WebSocket 핸들러에서 "
         "직접 호출 시 이벤트 루프가 블로킹되는 문제가 발생할 수 있다."),
        ("개선방안 및 향후계획",
         "박준우 : VIDEO 모드를 IMAGE 모드로 전환하여 단일 프레임 분석에 최적화하였고, "
         "asyncio.to_thread로 CPU 바운드 작업을 별도 스레드로 위임하여 이벤트 루프 블로킹을 해결하였다. "
         "향후에는 안드로이드 클라이언트에서 카메라 프레임을 WebSocket으로 전송하여 실시간 자세 피드백을 받는 기능을 통합할 예정이다."),
    ]

    for i, (label, body) in enumerate(sections):
        p = content_cell.add_paragraph() if i > 0 else content_cell.paragraphs[0]
        run = p.add_run(f"[{label}]")
        set_run_font(run, size=11, bold=True)
        p2 = content_cell.add_paragraph()
        run2 = p2.add_run(f"  - {body}")
        set_run_font(run2, size=11)


def build_detail_section(doc):
    """2페이지 이후: 박준우 상세 파트"""
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("① 박준우")
    set_run_font(run, size=16, bold=True)

    # 1. 구현 내용
    add_heading(doc, "1. 구현 내용", size=14, bold=True)

    add_heading(doc, "1-1. 분석 로직 함수화 (test_pose_8.py)", size=12, bold=True)
    add_paragraph(doc,
        "기존 test_pose_7.py는 cv2.VideoCapture를 통해 웹캠에서 프레임을 받아 while 루프로 "
        "반복 처리하는 단일 스크립트 구조였다. 이를 모바일 안드로이드 환경에서 호출 가능하도록 "
        "재사용 가능한 모듈 형태로 리팩토링하였다.",
        size=10)
    add_paragraph(doc, "주요 변경 사항은 다음과 같다.", size=10)
    add_paragraph(doc,
        "  · 웹캠 관련 코드(cv2.VideoCapture, cv2.imshow, cv2.waitKey)를 모두 제거하였다.\n"
        "  · MediaPipe PoseLandmarker 초기화를 모듈 로드 시 1회만 수행하도록 모듈 레벨로 이동하였다.\n"
        "  · MediaPipe 실행 모드를 VIDEO에서 IMAGE 모드로 전환하여 단일 프레임 분석에 적합하게 수정하였다.\n"
        "  · 핵심 분석 로직을 analyze_pose(image) 함수로 캡슐화하였다.\n"
        "  · SquatCounter, FeedbackManager 등 stateful 클래스는 외부 호출자가 사용할 수 있도록 보존하였다.",
        size=10)

    add_paragraph(doc, "[analyze_pose 함수 구현 코드]", size=10, bold=True)
    add_code(doc,
"""# 모듈 레벨 1회 초기화 (재사용 가능)
_options = PoseLandmarkerOptions(
    base_options=BaseOptions(model_asset_path=model_path),
    running_mode=VisionRunningMode.IMAGE,
)
_landmarker = PoseLandmarker.create_from_options(_options)

def analyze_pose(image):
    \"\"\"
    BGR numpy 이미지 한 장을 입력받아 자세 분석 결과를 dict로 반환한다.
    Returns: {
        "posture":  "good" | "bad",
        "feedback": 자세 교정 메시지 (문자열),
        "angles":   { "left_knee", "right_knee", "left_hip", "right_hip" }
    }
    \"\"\"
    rgb_frame = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    mp_image  = mp.Image(image_format=mp.ImageFormat.SRGB, data=rgb_frame)
    result    = _landmarker.detect(mp_image)

    if not result.pose_landmarks:
        return {"posture": "bad", "feedback": "Person not detected", "angles": {}}

    landmarks = result.pose_landmarks[0]
    angle_lk = calculate_angle(xy(LEFT_HIP),  xy(LEFT_KNEE),  xy(LEFT_ANKLE))
    angle_rk = calculate_angle(xy(RIGHT_HIP), xy(RIGHT_KNEE), xy(RIGHT_ANKLE))
    angle_lh = calculate_angle(xy(LEFT_SHOULDER),  xy(LEFT_HIP),  xy(LEFT_KNEE))
    angle_rh = calculate_angle(xy(RIGHT_SHOULDER), xy(RIGHT_HIP), xy(RIGHT_KNEE))

    angles = {"left_knee": angle_lk, "right_knee": angle_rk,
              "left_hip":  angle_lh, "right_hip":  angle_rh}

    avg_knee = (angle_lk + angle_rk) / 2
    if   avg_knee < SQUAT_DOWN_THRESHOLD: stage = "DOWN"
    elif avg_knee > SQUAT_UP_THRESHOLD:   stage = "UP"
    else:                                  stage = "MID"

    is_normal, issues = True, []
    if stage == "DOWN":
        is_normal, issues = judge_squat_pose(landmarks, angle_lk, angle_rk)

    posture = "good" if is_normal else "bad"
    if issues:
        msgs = [FEEDBACK_MESSAGES.get(i["key"], i["label"]) for i in issues]
        feedback_msg = " | ".join(msgs)
    elif stage == "DOWN":
        feedback_msg = "Good squat form!"
    else:
        feedback_msg = "Stand by"

    return {"posture": posture, "feedback": feedback_msg, "angles": angles}""")

    add_heading(doc, "1-2. FastAPI 서버 구현 (pose_server.py - REST)", size=12, bold=True)
    add_paragraph(doc,
        "FastAPI를 사용하여 자세 분석 REST API 서버를 구현하였다. "
        "POST /analyze 엔드포인트가 base64로 인코딩된 이미지를 수신하고, "
        "이를 numpy 배열로 디코딩한 뒤 analyze_pose 함수를 호출하여 결과를 JSON으로 반환한다.",
        size=10)
    add_paragraph(doc, "구현의 주요 포인트는 다음과 같다.", size=10)
    add_paragraph(doc,
        "  · CORSMiddleware를 추가하여 모바일/웹 어디서든 접근이 가능하도록 하였다.\n"
        "  · Pydantic BaseModel로 요청/응답 스키마를 정의하여 자동 검증을 적용하였다.\n"
        "  · data URL 접두어(data:image/jpeg;base64,...)를 자동으로 제거하도록 처리하였다.\n"
        "  · base64 디코딩 실패, 이미지 형식 인식 실패, 분석 중 예외 등을 단계별로 분리하여 적절한 HTTP 상태 코드를 반환하도록 하였다.",
        size=10)

    add_paragraph(doc, "[POST /analyze 엔드포인트 구현 코드]", size=10, bold=True)
    add_code(doc,
"""app = FastAPI(title="ModuFlow Pose Analyzer", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class AnalyzeRequest(BaseModel):
    image: str = Field(..., description="base64 인코딩된 이미지 문자열")

class AnalyzeResponse(BaseModel):
    posture: str
    feedback: str
    angles: dict

def decode_base64_image(image_b64: str) -> np.ndarray:
    if not image_b64:
        raise HTTPException(400, "image 필드가 비어 있습니다.")
    if image_b64.startswith("data:"):
        image_b64 = image_b64.split(",", 1)[1]
    try:
        raw = base64.b64decode(image_b64, validate=True)
    except (binascii.Error, ValueError):
        raise HTTPException(400, "base64 디코딩에 실패했습니다.")
    try:
        pil_img = Image.open(io.BytesIO(raw)).convert("RGB")
    except (UnidentifiedImageError, OSError):
        raise HTTPException(400, "이미지 파일 형식을 인식할 수 없습니다.")
    rgb = np.array(pil_img)
    return cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)

@app.post("/analyze", response_model=AnalyzeResponse)
def analyze(req: AnalyzeRequest):
    image = decode_base64_image(req.image)
    try:
        result = analyze_pose(image)
    except Exception as e:
        raise HTTPException(500, f"자세 분석 중 오류: {e}")
    return result""")

    add_heading(doc, "1-3. WebSocket 엔드포인트 추가 (pose_server.py - /ws)", size=12, bold=True)
    add_paragraph(doc,
        "REST API는 매 요청마다 연결을 새로 수립하므로 실시간 영상 분석에는 오버헤드가 크다. "
        "이를 보완하기 위해 WebSocket 엔드포인트 /ws를 추가 구현하여, 단일 연결을 유지한 상태에서 "
        "프레임을 연속적으로 송수신할 수 있도록 하였다.",
        size=10)
    add_paragraph(doc, "구현 시 핵심적으로 고려한 사항은 다음과 같다.", size=10)
    add_paragraph(doc,
        "  · async/await 기반의 비동기 핸들러로 작성하여 단일 프로세스에서 다수의 클라이언트 처리가 가능하도록 하였다.\n"
        "  · analyze_pose는 MediaPipe 추론을 수행하는 CPU 바운드 동기 함수이므로, asyncio.to_thread()를 통해 별도 스레드로 위임하여 이벤트 루프 블로킹을 방지하였다.\n"
        "  · 메시지 단위의 예외 처리를 통해 개별 프레임 처리에 실패하더라도 연결은 유지되도록 하였다.\n"
        "  · WebSocketDisconnect 예외를 분리하여 정상 종료와 비정상 종료를 구분 처리하였다.",
        size=10)

    add_paragraph(doc, "[WebSocket 엔드포인트 구현 코드]", size=10, bold=True)
    add_code(doc,
"""def _process_frame_blocking(image_b64: str) -> dict:
    \"\"\"디코딩 + 자세 분석 (CPU bound). 별도 스레드에서 실행됨.\"\"\"
    image = _decode_image_safe(image_b64)
    return analyze_pose(image)

@app.websocket("/ws")
async def ws_analyze(websocket: WebSocket):
    await websocket.accept()
    client = f"{websocket.client.host}:{websocket.client.port}"
    logger.info("WS 연결: %s", client)

    try:
        while True:
            try:
                msg = await websocket.receive_json()
            except ValueError:
                await websocket.send_json({"type": "error", "message": "JSON 파싱 실패"})
                continue

            if msg.get("type") != "frame":
                await websocket.send_json({
                    "type": "error",
                    "message": f"지원하지 않는 type: {msg.get('type')}",
                })
                continue

            try:
                # CPU 바운드 작업을 스레드로 분리
                result = await asyncio.to_thread(
                    _process_frame_blocking, msg.get("image", "")
                )
            except ValueError as e:
                await websocket.send_json({"type": "error", "message": str(e)})
                continue
            except Exception as e:
                logger.exception("자세 분석 중 예외")
                await websocket.send_json({
                    "type": "error",
                    "message": f"자세 분석 중 오류: {e}",
                })
                continue

            await websocket.send_json({
                "type":     "result",
                "posture":  result["posture"],
                "feedback": result["feedback"],
                "angles":   result["angles"],
            })

    except WebSocketDisconnect:
        logger.info("WS 연결 해제: %s", client)
    except Exception as e:
        logger.exception("WS 처리 중 예기치 않은 예외: %s", e)
        try:
            await websocket.close(code=1011)
        except Exception:
            pass""")

    # 2. 실행 결과 확인
    add_heading(doc, "2. 실행 결과 확인", size=14, bold=True)

    add_heading(doc, "2-1. 설치한 라이브러리", size=12, bold=True)
    add_paragraph(doc,
        "본 8주차 작업을 위해 다음의 Python 라이브러리를 추가로 설치하였다.",
        size=10)
    add_code(doc,
"""pip install fastapi uvicorn pillow python-multipart""")
    add_paragraph(doc,
        "  · fastapi : 고성능 비동기 웹 프레임워크 (REST + WebSocket 지원)\n"
        "  · uvicorn : FastAPI 애플리케이션 실행을 위한 ASGI 서버\n"
        "  · pillow  : base64 디코딩된 이미지의 형식 판별 및 RGB 변환\n"
        "  · python-multipart : FastAPI의 form 데이터 처리 의존성",
        size=10)

    add_heading(doc, "2-2. 서버 실행 명령어", size=12, bold=True)
    add_paragraph(doc,
        "MediaPipe 모델 파일이 상대경로로 지정되어 있어 ai_pose_project/src 디렉토리에서 실행한다.",
        size=10)
    add_code(doc,
"""cd ai_pose_project/src
uvicorn pose_server:app --host 0.0.0.0 --port 8000 --reload""")

    add_heading(doc, "2-3. 메시지 프로토콜", size=12, bold=True)
    add_paragraph(doc, "[REST: POST /analyze]", size=10, bold=True)
    add_code(doc,
"""// 요청
{ "image": "<base64 인코딩된 이미지>" }

// 응답
{
  "posture":  "good",
  "feedback": "Good squat form!",
  "angles":   {
    "left_knee":  95.2,
    "right_knee": 96.7,
    "left_hip":   100.1,
    "right_hip":  99.8
  }
}""")

    add_paragraph(doc, "[WebSocket: /ws]", size=10, bold=True)
    add_code(doc,
"""// 클라이언트 → 서버
{ "type": "frame", "image": "<base64>" }

// 서버 → 클라이언트 (성공)
{
  "type":     "result",
  "posture":  "bad",
  "feedback": "Tip: Push LEFT knee back, align over ankle",
  "angles":   { "left_knee": 88.5, "right_knee": 110.2, ... }
}

// 서버 → 클라이언트 (실패)
{ "type": "error", "message": "base64 디코딩에 실패했습니다." }""")

    # 3. 문제점 및 해결 방안
    add_heading(doc, "3. 문제점 및 해결 방안", size=14, bold=True)

    add_heading(doc, "3-1. MediaPipe 실행 모드 부적합 문제", size=12, bold=True)
    add_paragraph(doc,
        "기존 test_pose_7.py는 MediaPipe의 VIDEO 모드를 사용하여 cv2.CAP_PROP_POS_MSEC로 "
        "단조 증가하는 timestamp를 전달하는 구조였다. 그러나 외부 호출자가 임의의 단일 프레임을 "
        "분석하는 함수에서는 이 timestamp를 보장하기가 어려운 문제가 있었다.",
        size=10)
    add_paragraph(doc,
        "이를 해결하기 위해 PoseLandmarker의 실행 모드를 VisionRunningMode.IMAGE로 변경하고, "
        "detect_for_video(image, timestamp) 대신 detect(image)를 호출하도록 수정하였다. "
        "IMAGE 모드는 각 프레임을 독립적으로 처리하므로 함수형 인터페이스에 적합하다.",
        size=10)

    add_heading(doc, "3-2. 동기 함수의 이벤트 루프 블로킹 문제", size=12, bold=True)
    add_paragraph(doc,
        "analyze_pose는 MediaPipe 추론을 수행하는 동기 함수로, 한 프레임당 수십~수백 ms의 시간이 소요된다. "
        "이를 비동기 WebSocket 핸들러에서 직접 호출할 경우 해당 시간 동안 이벤트 루프 전체가 블로킹되어 "
        "다른 클라이언트 요청을 처리할 수 없는 문제가 발생한다.",
        size=10)
    add_paragraph(doc,
        "이를 해결하기 위해 Python 3.9에서 도입된 asyncio.to_thread()를 사용하여 CPU 바운드 작업을 "
        "별도 스레드에서 실행하도록 하였다. 이를 통해 메인 이벤트 루프는 블로킹되지 않고 다른 "
        "WebSocket 연결의 메시지 송수신을 계속 처리할 수 있게 되었다.",
        size=10)

    add_heading(doc, "3-3. 잘못된 입력에 대한 연결 종료 문제", size=12, bold=True)
    add_paragraph(doc,
        "WebSocket 핸들러 내부에서 예외 처리를 단일 try-except로 작성할 경우, 한 프레임의 base64 디코딩 "
        "실패만으로 전체 연결이 종료되는 문제가 있었다. 이는 실시간 영상 스트리밍 중 일시적인 노이즈로 "
        "인한 디코딩 실패 시에도 클라이언트가 재연결을 해야 하는 비효율을 초래한다.",
        size=10)
    add_paragraph(doc,
        "이를 해결하기 위해 예외 처리를 두 단계로 분리하였다. 외부 try-except는 WebSocketDisconnect 등 "
        "연결 자체의 종료를 처리하고, 내부 try-except는 개별 메시지의 처리 실패를 처리하여 "
        "에러 메시지만 전송하고 루프를 계속하도록 구현하였다.",
        size=10)

    # 4. 향후 계획
    add_heading(doc, "4. 향후 계획", size=14, bold=True)
    add_paragraph(doc,
        "8주차에 구축한 자세 분석 서버를 기반으로 향후에는 다음의 작업을 수행할 예정이다.",
        size=10)
    add_paragraph(doc,
        "  ① 안드로이드 클라이언트 앱에서 카메라 프레임을 캡처하여 WebSocket으로 서버에 전송하고, "
        "수신한 분석 결과를 화면에 시각화하는 모바일 클라이언트를 구현한다.\n"
        "  ② 현재는 단일 프레임 단위의 stateless 분석만 지원하므로, 세션 단위로 SquatCounter와 "
        "FeedbackManager의 상태를 유지하는 stateful WebSocket 핸들러를 추가 구현한다.\n"
        "  ③ MediaPipe 추론 시간이 모바일 환경의 실시간 요구사항(약 30fps)에 부합하는지 측정하고, "
        "필요 시 프레임 스킵 또는 더 가벼운 모델로의 교체를 검토한다.\n"
        "  ④ 7주차에 구현한 Spring 백엔드와의 연동을 유지하면서, FastAPI 분석 서버와 Spring 비즈니스 서버를 "
        "어떻게 분리·통합할지에 대한 아키텍처 결정을 내린다.",
        size=10)


def main():
    doc = Document()

    # 기본 폰트 설정
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)

    build_summary_table(doc)
    build_detail_section(doc)

    today = datetime.now()
    out_dir = r"c:\Users\junwoo\Desktop\4학년1학기\종합설계프로젝트\연구노트"
    os.makedirs(out_dir, exist_ok=True)
    filename = f"연구노트_{today.year}_{today.month:02d}_{today.day:02d}.docx"
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)
    print(f"[OK] 연구노트 저장: {out_path}")


if __name__ == "__main__":
    main()
