"""
9주차 종합설계프로젝트 연구노트 생성 스크립트

8주차에 구현한 분석 서버의 검증 및 실시간 통합 테스트 내용을 기록한다.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "맑은 고딕"


def set_run_font(run, name=FONT_NAME, size=10, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_paragraph(doc, text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_code(doc, code_text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return p


def set_cell_text(cell, text, size=10, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_summary_table(doc):
    add_heading(doc, "종합설계프로젝트 연구노트 (9주차)", size=16, bold=True)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(13)

    set_cell_text(table.cell(0, 0), "연구일시", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "2026년 04월 30일")

    set_cell_text(table.cell(1, 0), "연구장소", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 1), "경성대학교")

    set_cell_text(table.cell(2, 0), "수행자", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 1), "박준우, 이경민, 임용완, 정주영")

    set_cell_text(table.cell(3, 0), "연구내용", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    content_cell = table.cell(3, 1)
    content_cell.text = ""

    sections = [
        ("목적",
         "박준우 : 8주차에 구현한 analyze_pose 모듈과 FastAPI 서버(REST/WebSocket)의 동작을 검증하고, "
         "실시간 웹캠 스트리밍 클라이언트를 구현하여 끝-끝(end-to-end) 통합 테스트를 수행하였다."),
        ("연구방법",
         "박준우 : 더미 이미지 생성기와 웹캠 캡처 헬퍼를 구현하여 회귀 테스트 데이터셋을 구축하였고, "
         "단계별 테스트 스크립트(test_analyze.py, test_client.py, live_client.py)를 작성하였다. "
         "uvicorn 서버를 백그라운드로 실행한 상태에서 requests 및 websockets 라이브러리로 "
         "REST와 WebSocket 양쪽 엔드포인트를 검증하였다."),
        ("도출결과",
         "박준우 : 모듈 단독 / REST API / WebSocket / 실시간 웹캠 스트리밍의 4단계 검증을 모두 통과하였다. "
         "특히 실시간 스트리밍에서 단일 WebSocket 연결로 604 프레임을 처리하고 에러 0건을 달성하여, "
         "asyncio.to_thread 기반의 비동기 처리가 안정적으로 동작함을 확인하였다."),
        ("문제점 분석",
         "박준우 : analyze_pose 모듈을 외부에서 import할 경우 모델 파일을 찾지 못하는 상대경로 의존 문제가 있었으며, "
         "백그라운드로 실행한 uvicorn 서버를 깔끔하게 종료하기 위한 프로세스 관리 절차도 필요하였다."),
        ("개선방안 및 향후계획",
         "박준우 : 모델 경로를 __file__ 기준 절대경로로 변경하여 호출자의 CWD에 무관하게 동작하도록 수정하였다. "
         "다음 주차에는 안드로이드 클라이언트 앱(Kotlin + CameraX + OkHttp WebSocket)을 구현하여 "
         "실제 모바일 환경의 latency 및 FPS를 측정하고, 세션 단위의 stateful 핸들러를 추가할 예정이다."),
    ]

    for i, (label, body) in enumerate(sections):
        p = content_cell.add_paragraph() if i > 0 else content_cell.paragraphs[0]
        run = p.add_run(f"[{label}]")
        set_run_font(run, size=11, bold=True)
        p2 = content_cell.add_paragraph()
        run2 = p2.add_run(f"  - {body}")
        set_run_font(run2, size=11)


def build_detail_section(doc):
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("① 박준우")
    set_run_font(run, size=16, bold=True)

    # 1. 구현 내용
    add_heading(doc, "1. 구현 내용", size=14, bold=True)

    # 1-1
    add_heading(doc, "1-1. 테스트 이미지 데이터셋 구축", size=12, bold=True)
    add_paragraph(doc,
        "자세 분석 모듈의 회귀 테스트를 위해 ai_pose_project/test_images/ 디렉토리를 생성하고 "
        "두 종류의 이미지 데이터셋을 구축하였다. 첫째는 사람이 없는 더미 이미지 3종으로 "
        "\"Person not detected\" 응답을 검증하기 위한 것이고, 둘째는 실제 웹캠으로 캡처한 "
        "사람 사진 3장으로 정상 분석 응답을 검증하기 위한 것이다.",
        size=10)
    add_paragraph(doc, "[더미 이미지 생성기 _generate.py]", size=10, bold=True)
    add_code(doc,
"""def make_blank_black():
    img = Image.new("RGB", (640, 480), color=(0, 0, 0))
    img.save(os.path.join(OUT_DIR, "blank_black.jpg"), quality=85)

def make_noise():
    arr = np.random.randint(0, 256, (480, 640, 3), dtype=np.uint8)
    Image.fromarray(arr).save(os.path.join(OUT_DIR, "noise.jpg"), quality=85)

def make_stick_figure():
    img = Image.new("RGB", (640, 480), color=(230, 230, 230))
    d = ImageDraw.Draw(img)
    d.ellipse((300, 70, 340, 115), outline=(0, 0, 0), width=3)   # 머리
    d.line((320, 115, 320, 280), fill=(0, 0, 0), width=3)         # 몸통
    d.line((320, 160, 240, 240), fill=(0, 0, 0), width=3)         # 좌팔
    d.line((320, 160, 400, 240), fill=(0, 0, 0), width=3)         # 우팔
    d.line((320, 280, 270, 410), fill=(0, 0, 0), width=3)         # 좌다리
    d.line((320, 280, 370, 410), fill=(0, 0, 0), width=3)         # 우다리
    img.save(os.path.join(OUT_DIR, "stick_figure.jpg"), quality=85)""")

    add_paragraph(doc, "[웹캠 캡처 헬퍼 _capture.py]", size=10, bold=True)
    add_paragraph(doc,
        "OpenCV의 VideoCapture로 웹캠 영상을 표시하고, SPACE 키 입력 시 현재 프레임을 "
        "person_capture_YYYYMMDD_HHMMSS.jpg 형태의 파일로 저장하는 헬퍼를 구현하였다. "
        "ESC 키로 종료된다. 본 헬퍼를 사용하여 서기 자세 2장과 잘못된 스쿼트 자세 1장을 캡처하였다.",
        size=10)

    # 1-2
    add_heading(doc, "1-2. 모듈 단독 테스트 (test_analyze.py)", size=12, bold=True)
    add_paragraph(doc,
        "test_images/ 디렉토리의 모든 이미지에 대해 analyze_pose() 함수를 호출하고 "
        "반환 dict를 JSON 형태로 출력하는 테스트 스크립트를 작성하였다. "
        "test_pose_8 모듈이 import 시점에 모델 파일을 로드하므로, 해당 시점에 CWD를 "
        "src 디렉토리로 변경하는 처리를 포함하였다.",
        size=10)
    add_code(doc,
"""HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)

import cv2
from test_pose_8 import analyze_pose

TEST_DIR = os.path.normpath(os.path.join(HERE, "..", "test_images"))

def main():
    files = sorted(
        f for f in os.listdir(TEST_DIR)
        if f.lower().endswith((".jpg", ".jpeg", ".png"))
        and not f.startswith("_")
    )
    for f in files:
        path = os.path.join(TEST_DIR, f)
        img = cv2.imread(path)
        result = analyze_pose(img)
        print(f"파일: {f}   ({img.shape[1]}x{img.shape[0]})")
        print(json.dumps(result, ensure_ascii=False, indent=2))""")

    # 1-3
    add_heading(doc, "1-3. REST + WebSocket 통합 테스트 (test_client.py)", size=12, bold=True)
    add_paragraph(doc,
        "FastAPI 서버에 대한 통합 테스트 스크립트를 작성하였다. 헬스체크 → REST 정상 케이스 → "
        "REST 에러 케이스 → WebSocket 정상 케이스 → WebSocket 에러 케이스 → 연결 유지 검증의 "
        "순서로 단계별로 검증한다. 동일한 이미지에 대해 REST와 WebSocket이 동일한 응답을 "
        "반환하는지 확인함으로써, 두 엔드포인트가 같은 analyze_pose 함수를 호출함을 검증하였다.",
        size=10)
    add_code(doc,
"""async def test_websocket():
    files = list_images()
    async with websockets.connect(WS_URL) as ws:
        # 1) 정상 케이스: 여러 프레임 연속 전송
        for f in files:
            b64 = encode_b64(os.path.join(TEST_DIR, f))
            await ws.send(json.dumps({"type": "frame", "image": b64}))
            data = json.loads(await ws.recv())
            print(f"[{f}]  posture={data.get('posture')}")

        # 2) 에러: 지원하지 않는 type
        await ws.send(json.dumps({"type": "wrong", "image": ""}))
        print(await ws.recv())

        # 3) 에러: 잘못된 base64
        await ws.send(json.dumps({"type": "frame", "image": "!!!not_base64"}))
        print(await ws.recv())

        # 4) 연결 유지 검증
        b64 = encode_b64(os.path.join(TEST_DIR, files[0]))
        await ws.send(json.dumps({"type": "frame", "image": b64}))
        data = json.loads(await ws.recv())
        print(f"  type={data['type']}  → 연결 유지 OK")""")

    # 1-4
    add_heading(doc, "1-4. 실시간 웹캠 스트리밍 클라이언트 (live_client.py)", size=12, bold=True)
    add_paragraph(doc,
        "웹캠 영상을 실시간으로 서버에 전송하고 분석 결과를 화면에 오버레이하는 비동기 클라이언트를 구현하였다. "
        "각 프레임을 JPEG으로 압축(quality 70)한 뒤 base64로 인코딩하여 WebSocket으로 송신하고, "
        "수신한 결과를 OpenCV로 화면에 표시한다. FPS와 통신 latency를 EMA(지수이동평균) 방식으로 "
        "계산하여 함께 표시한다.",
        size=10)
    add_code(doc,
"""async def run():
    cap = cv2.VideoCapture(0)
    async with websockets.connect(WS_URL, max_size=4*1024*1024) as ws:
        fps, last_t = 0.0, time.time()
        while True:
            ret, frame = cap.read()
            if not ret: break

            ok, buf = cv2.imencode(".jpg", frame,
                                   [cv2.IMWRITE_JPEG_QUALITY, JPEG_QUALITY])
            b64 = base64.b64encode(buf.tobytes()).decode("ascii")

            t0 = time.time()
            await ws.send(json.dumps({"type": "frame", "image": b64}))
            data = json.loads(await ws.recv())
            latency_ms = (time.time() - t0) * 1000

            now = time.time()
            inst_fps = 1.0 / (now - last_t)
            fps = 0.85 * fps + 0.15 * inst_fps if fps else inst_fps
            last_t = now

            if data.get("type") == "result":
                draw_overlay(frame, data, fps, latency_ms)
            cv2.imshow("Live Pose Analysis", frame)
            if cv2.waitKey(1) & 0xFF == 27:
                break""")

    # 2. 실행 결과 확인
    add_heading(doc, "2. 실행 결과 확인", size=14, bold=True)

    add_heading(doc, "2-1. 모듈 단독 테스트 결과", size=12, bold=True)
    add_paragraph(doc,
        "총 6개 이미지에 대해 analyze_pose()를 호출한 결과는 모두 의도한 분류와 일치하였다.",
        size=10)
    add_code(doc,
"""파일명                                  posture  비고
─────────────────────────────────────────────────────────────────
blank_black.jpg                         bad      Person not detected
noise.jpg                               bad      Person not detected
stick_figure.jpg                        bad      Person not detected
person_capture_20260430_143528.jpg      good     무릎 176/175°  stage UP
person_capture_20260430_143532.jpg      good     무릎 179/178°  stage UP
person_capture_20260430_143536.jpg      bad      무릎 86/97°  stage DOWN
                                                 → "L knee back | R knee back | Lift chest" """)

    add_heading(doc, "2-2. REST + WebSocket 통합 테스트 결과", size=12, bold=True)
    add_paragraph(doc,
        "REST 엔드포인트는 정상 6건 + 에러 2건 모두 통과하였고, WebSocket 엔드포인트는 "
        "정상 6건 + 에러 3건 + 연결 유지 검증을 모두 통과하였다. 같은 이미지에 대한 "
        "REST와 WebSocket 응답이 100% 일치함을 확인하였다.",
        size=10)
    add_code(doc,
"""[REST]
  blank_black.jpg                       HTTP 200  posture=bad
  noise.jpg                             HTTP 200  posture=bad
  person_capture_20260430_143528.jpg    HTTP 200  posture=good
  person_capture_20260430_143532.jpg    HTTP 200  posture=good
  person_capture_20260430_143536.jpg    HTTP 200  posture=bad
  stick_figure.jpg                      HTTP 200  posture=bad
  잘못된 base64                          HTTP 400  base64 디코딩 실패
  빈 image 필드                          HTTP 400  image 필드 비어있음

[WebSocket]
  6장 정상 응답 (posture/feedback/angles 모두 REST와 일치)
  지원하지 않는 type        →  type=error  message=지원하지 않는 type: wrong
  잘못된 base64            →  type=error  message=base64 디코딩 실패
  에러 후 정상 요청         →  type=result  → 연결 유지 OK""")

    add_heading(doc, "2-3. 실시간 웹캠 스트리밍 결과", size=12, bold=True)
    add_paragraph(doc,
        "단일 WebSocket 연결로 약 1분간 실시간 스트리밍을 진행한 결과, 총 604개의 프레임이 "
        "처리되고 에러는 단 한 건도 발생하지 않았다. 이는 asyncio.to_thread를 통한 CPU 바운드 "
        "분리, JPEG 압축 + base64 인코딩 파이프라인, 그리고 WebSocket의 단일 연결 유지가 "
        "안정적으로 동작함을 입증한다.",
        size=10)
    add_code(doc,
"""[연결 시도] ws://127.0.0.1:8000/ws
[연결 OK]  ESC로 종료
[종료] 수신 결과 604건 / 에러 0건""")

    # 3. 문제점 및 해결 방안
    add_heading(doc, "3. 문제점 및 해결 방안", size=14, bold=True)

    add_heading(doc, "3-1. 모델 파일 경로의 호출자 의존 문제", size=12, bold=True)
    add_paragraph(doc,
        "test_pose_8.py가 모델 파일을 'pose_landmarker_lite.task'라는 상대경로로 로드하도록 "
        "작성되어 있어, 외부 모듈에서 import할 경우 호출자의 CWD에 따라 모델 파일을 찾지 못하는 "
        "FileNotFoundError가 발생하였다. 특히 test_analyze.py에서 cwd를 src로 변경했음에도 "
        "모델 파일이 src의 부모 디렉토리에 있어 실패하였다.",
        size=10)
    add_paragraph(doc,
        "이를 해결하기 위해 모델 경로를 __file__ 기준의 절대경로로 변경하여 "
        "어느 디렉토리에서 import하든 동일하게 동작하도록 수정하였다. "
        "이 수정으로 test_analyze.py, pose_server.py(REST/WebSocket) 모두 정상 동작하게 되었다.",
        size=10)
    add_code(doc,
"""# 수정 전
model_path = "pose_landmarker_lite.task"

# 수정 후
model_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..",
    "pose_landmarker_lite.task",
)""")

    add_heading(doc, "3-2. 백그라운드 서버 프로세스 관리", size=12, bold=True)
    add_paragraph(doc,
        "통합 테스트를 위해 uvicorn 서버를 백그라운드로 실행한 뒤 클라이언트 스크립트를 "
        "실행하는 절차에서, 테스트 종료 후 서버 프로세스를 정리하지 않으면 8000 포트가 "
        "계속 점유되어 다음 실행에 영향을 주는 문제가 있었다.",
        size=10)
    add_paragraph(doc,
        "이를 해결하기 위해 백그라운드 task의 ID를 추적하고, 테스트 완료 후 명시적으로 "
        "종료하는 절차를 정착시켰다. 또한 서버 시작 직후 곧바로 클라이언트를 실행하면 "
        "서버 미준비 상태로 인해 connection refused가 발생하므로, /health 엔드포인트를 "
        "1초 간격으로 폴링하여 서버 준비 완료 후 테스트가 시작되도록 처리하였다.",
        size=10)

    add_heading(doc, "3-3. 한글 콘솔 출력 인코딩 문제", size=12, bold=True)
    add_paragraph(doc,
        "Windows 명령 프롬프트의 기본 코드 페이지가 cp949이고 Python 3.10의 기본 stdout "
        "인코딩과 충돌하여, 테스트 스크립트의 한글 출력이 깨져 보이는 문제가 있었다. "
        "다만 결과 dict 자체는 정상적으로 처리되어 검증에는 영향이 없었으며, "
        "향후 PYTHONIOENCODING=utf-8 환경 변수 설정 또는 chcp 65001을 통한 콘솔 코드 페이지 "
        "변경으로 보완할 예정이다.",
        size=10)

    # 4. 향후 계획
    add_heading(doc, "4. 향후 계획", size=14, bold=True)
    add_paragraph(doc,
        "9주차에 완료한 서버 검증을 기반으로 다음 주차에는 모바일 클라이언트 구현으로 진행한다.",
        size=10)
    add_paragraph(doc,
        "  ① Kotlin + CameraX + OkHttp WebSocket을 사용한 안드로이드 클라이언트 앱을 구현한다. "
        "live_client.py와 동일한 메시지 프로토콜을 따르므로 즉시 연동 가능할 것으로 예상된다.\n"
        "  ② 안드로이드 단말의 카메라 프레임을 JPEG으로 압축하여 WebSocket으로 송신하고, "
        "Jetpack Compose로 분석 결과를 화면에 오버레이한다.\n"
        "  ③ 실제 모바일 환경에서의 latency, FPS, 배터리 소모를 측정하여 "
        "JPEG 압축률, 프레임 스킵 등의 최적화 파라미터를 결정한다.\n"
        "  ④ 현재 stateless인 분석 함수를 보완하여, 세션 단위로 SquatCounter와 FeedbackManager의 "
        "상태를 유지하는 stateful WebSocket 핸들러를 추가 구현한다.\n"
        "  ⑤ 7주차에 구축한 Spring 백엔드와의 연동 아키텍처를 정리하고, "
        "FastAPI 분석 서버와 Spring 비즈니스 서버의 책임 분리 방안을 확정한다.",
        size=10)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)

    build_summary_table(doc)
    build_detail_section(doc)

    today = datetime.now()
    out_dir = r"c:\Users\junwoo\Desktop\4학년1학기\종합설계프로젝트\연구노트"
    os.makedirs(out_dir, exist_ok=True)
    filename = f"연구노트_9주차_{today.year}_{today.month:02d}_{today.day:02d}.docx"
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)
    print(f"[OK] 9주차 연구노트 저장: {out_path}")


if __name__ == "__main__":
    main()
