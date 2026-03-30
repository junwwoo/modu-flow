"""5주차 연구노트 생성 스크립트"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ──────────────────────────────────────────────────────────────
# 기본 스타일 설정
# ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ──────────────────────────────────────────────────────────────
# 여백 설정
# ──────────────────────────────────────────────────────────────
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_font(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """셀 내부 텍스트 서식 설정"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.bold = bold
    return run


def set_cell_shading(cell, color="D9E2F3"):
    """셀 배경색 설정"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_styled_paragraph(doc, text, size=10, bold=False, space_after=6, space_before=0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """스타일이 적용된 문단 추가"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.bold = bold
    return p


def add_code_block(doc, code_text):
    """코드 블록을 문서에 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 코드 배경 표현을 위해 들여쓰기 활용
    run = p.add_run(code_text)
    run.font.size = Pt(8.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p


# ══════════════════════════════════════════════════════════════
# 1페이지: 요약 페이지
# ══════════════════════════════════════════════════════════════
add_styled_paragraph(doc, "종합설계프로젝트 연구노트", size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# 기본 정보 표
info_table = doc.add_table(rows=3, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

# 열 너비 설정
for row in info_table.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(13)

info_data = [
    ("연구일시", "2026년 03월 30일"),
    ("연구장소", "경성대학교"),
    ("수행자",   "박준우, 이경민, 임용완, 정주영"),
]

for i, (label, value) in enumerate(info_data):
    set_cell_shading(info_table.rows[i].cells[0], "D9E2F3")
    set_cell_font(info_table.rows[i].cells[0], label, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_font(info_table.rows[i].cells[1], value, size=11)

add_styled_paragraph(doc, "", size=6, space_after=6)

# ──────────────────────────────────────────────────────────────
# 연구내용 섹션 (요약)
# ──────────────────────────────────────────────────────────────
add_styled_paragraph(doc, "■ 연구내용", size=13, bold=True, space_before=10, space_after=8)

summary_items = [
    ("목적",
     "5주차 과제로서, 스쿼트 자세 판별 시스템에 실시간 피드백 기능을 추가하여 "
     "사용자가 운동 중 즉각적인 자세 교정 가이드를 받을 수 있도록 구현하는 것을 목적으로 하였다."),

    ("연구방법",
     "Python, OpenCV, MediaPipe PoseLandmarker, NumPy를 활용하였다. "
     "4주차에서 구현한 스쿼트 자세 판별 로직을 기반으로, FeedbackManager 클래스를 새로 설계하여 "
     "문제 유형별 교정 메시지 매핑, 시간 기반 메시지 표시, rep 단위 요약, 세션 통계 기능을 구현하였다."),

    ("도출결과",
     "test_pose_5.py 파일을 생성하여 실시간 피드백 시스템을 완성하였다. "
     "무릎 전방 이탈, 상체 과도 전경, 좌우 비대칭 등의 문제가 감지되면 "
     "화면 우측에 구체적인 교정 가이드 메시지가 3초간 표시되며, "
     "1회 스쿼트 완료 시 해당 rep의 자세 평가 요약이 출력되고, "
     "프로그램 종료 시 전체 세션 통계가 터미널에 출력되도록 하였다."),

    ("문제점 분석",
     "OpenCV의 putText 함수가 한글 렌더링을 지원하지 않아 피드백 메시지를 영어로 작성하였다. "
     "향후 PIL(Pillow) 라이브러리를 연동하여 한글 표시를 지원할 수 있을 것으로 판단된다."),

    ("개선방안 및 향후계획",
     "6주차에서는 모바일(안드로이드) 환경으로의 전환을 준비하고, "
     "다양한 운동(런지, 플랭크 등)에 대한 자세 판별 및 피드백 기능을 확장할 계획이다. "
     "또한 피드백 메시지의 한글화 및 음성 피드백 기능 추가를 검토할 예정이다."),
]

for label, content in summary_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # "Ÿ 박준우 : " 접두어
    run_prefix = p.add_run(f"Ÿ 박준우 : {label} - ")
    run_prefix.font.size = Pt(11)
    run_prefix.font.name = '맑은 고딕'
    run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run_prefix.bold = True

    run_body = p.add_run(content)
    run_body.font.size = Pt(11)
    run_body.font.name = '맑은 고딕'
    run_body._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ══════════════════════════════════════════════════════════════
# 2페이지: 상세 내용 (박준우 파트)
# ══════════════════════════════════════════════════════════════
doc.add_page_break()

add_styled_paragraph(doc, "① 박준우", size=14, bold=True, space_after=10)

# ── 1. 개요 ──
add_styled_paragraph(doc, "1. 5주차 작업 개요", size=14, bold=True, space_before=8, space_after=6)

add_styled_paragraph(doc,
    "5주차에서는 4주차까지 구현한 스쿼트 자세 판별 시스템에 실시간 피드백 기능을 추가하였다. "
    "기존에는 자세의 정상/비정상 여부만 표시하였으나, 5주차에서는 구체적인 교정 메시지를 "
    "화면에 표시하여 사용자가 즉각적으로 자세를 교정할 수 있도록 개선하였다.",
    size=10, space_after=6)

# ── 1-1. 추가된 기능 ──
add_styled_paragraph(doc, "1-1. 추가된 주요 기능", size=12, bold=True, space_before=6, space_after=4)

features = [
    "실시간 교정 메시지 표시: 문제 감지 시 구체적인 교정 가이드를 3초간 화면 우측에 표시",
    "문제 유형별 피드백 매핑: 무릎 전방 이탈, 상체 전경, 좌우 비대칭 각각에 대한 맞춤 메시지 제공",
    "Rep 단위 요약: 1회 스쿼트 완료 시 해당 rep의 자세 평가(Good form / 문제 요약) 표시",
    "세션 통계: 화면 하단에 전체 세션 동안의 문제 발생 횟수 누적 표시",
    "종료 시 세션 요약: ESC로 종료하면 터미널에 총 횟수와 문제 통계 출력",
]
for feat in features:
    p = doc.add_paragraph(feat, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ── 2. 핵심 구현 내용 ──
add_styled_paragraph(doc, "2. 핵심 구현 내용", size=14, bold=True, space_before=10, space_after=6)

# ── 2-1. 피드백 메시지 매핑 ──
add_styled_paragraph(doc, "2-1. 피드백 메시지 매핑", size=12, bold=True, space_before=6, space_after=4)

add_styled_paragraph(doc,
    "각 문제 유형에 대해 사용자가 이해할 수 있는 구체적인 교정 가이드 메시지를 매핑하였다. "
    "문제가 감지되면 해당 메시지가 화면 우측 패널에 표시된다.",
    size=10, space_after=4)

code_feedback_map = '''FEEDBACK_MESSAGES = {
    "left_knee_forward":  "Tip: Push LEFT knee back, align over ankle",
    "right_knee_forward": "Tip: Push RIGHT knee back, align over ankle",
    "trunk_lean":         "Tip: Lift your chest, keep torso upright",
    "knee_asymmetry":     "Tip: Balance both knees evenly",
}'''
add_code_block(doc, code_feedback_map)

# ── 2-2. FeedbackManager 클래스 ──
add_styled_paragraph(doc, "2-2. FeedbackManager 클래스", size=12, bold=True, space_before=8, space_after=4)

add_styled_paragraph(doc,
    "피드백 시스템의 핵심인 FeedbackManager 클래스를 설계하였다. "
    "이 클래스는 감지된 문제를 누적 추적하고, 시간 기반으로 피드백 메시지의 표시/만료를 관리하며, "
    "rep 단위 요약과 세션 통계를 집계하는 역할을 수행한다.",
    size=10, space_after=4)

code_feedback_mgr = '''class FeedbackManager:
    def __init__(self):
        self.active_feedbacks = {}      # {key: (message, expire_time)}
        self.current_rep_issues = set() # 현재 rep 문제 key 집합
        self.session_stats = {}         # {key: count} 전체 세션 통계
        self.rep_summary = ""
        self.rep_summary_expire = 0.0

    def update(self, issues, stage):
        """매 프레임 호출: 감지된 문제를 반영하여 피드백 상태를 갱신"""
        now = time.time()
        for issue in issues:
            key = issue["key"]
            self.current_rep_issues.add(key)
            self.session_stats[key] = self.session_stats.get(key, 0) + 1
            msg = FEEDBACK_MESSAGES.get(key, "")
            if msg:
                self.active_feedbacks[key] = (msg, now + FEEDBACK_DISPLAY_SEC)
        # 만료된 피드백 제거
        expired = [k for k, (_, t) in self.active_feedbacks.items() if now > t]
        for k in expired:
            del self.active_feedbacks[k]

    def on_rep_complete(self, rep_count):
        """1회 스쿼트 완료 시 rep 요약 생성"""
        if self.current_rep_issues:
            labels = [FEEDBACK_MESSAGES[k].replace("Tip: ", "")
                      for k in self.current_rep_issues if k in FEEDBACK_MESSAGES]
            self.rep_summary = f"Rep {rep_count}: {', '.join(labels)}"
        else:
            self.rep_summary = f"Rep {rep_count}: Good form!"
        self.rep_summary_expire = time.time() + FEEDBACK_DISPLAY_SEC
        self.current_rep_issues.clear()'''
add_code_block(doc, code_feedback_mgr)

# ── 2-3. 화면 표시 ──
add_styled_paragraph(doc, "2-3. 피드백 화면 표시 함수", size=12, bold=True, space_before=8, space_after=4)

add_styled_paragraph(doc,
    "화면 우측에 피드백 패널을 오버레이로 표시하는 draw_feedback_panel 함수를 구현하였다. "
    "활성 피드백 메시지와 rep 요약 메시지를 동적으로 표시하며, "
    "패널 크기는 메시지 수에 따라 자동 조절된다.",
    size=10, space_after=4)

code_draw_feedback = '''def draw_feedback_panel(frame, feedback_msgs, rep_summary):
    """화면 우측에 교정 피드백 메시지를 표시"""
    h, w = frame.shape[:2]
    if not feedback_msgs and not rep_summary:
        return
    panel_x = w - 420
    panel_y = 10
    panel_h = 30 + len(feedback_msgs) * 30 + (35 if rep_summary else 0)
    cv2.rectangle(frame, (panel_x, panel_y), (w - 10, panel_y + panel_h),
                  (40, 40, 40), -1)
    cv2.rectangle(frame, (panel_x, panel_y), (w - 10, panel_y + panel_h),
                  (0, 200, 255), 2)
    cv2.putText(frame, "FEEDBACK",
                (panel_x + 10, panel_y + 22),
                cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 200, 255), 2)
    for i, msg in enumerate(feedback_msgs):
        y_pos = panel_y + 50 + i * 30
        cv2.putText(frame, msg, (panel_x + 10, y_pos),
                    cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 1)'''
add_code_block(doc, code_draw_feedback)

# ── 2-4. 메인 루프 통합 ──
add_styled_paragraph(doc, "2-4. 메인 루프 통합", size=12, bold=True, space_before=8, space_after=4)

add_styled_paragraph(doc,
    "기존 메인 루프에 FeedbackManager를 통합하여, 매 프레임마다 피드백을 갱신하고 "
    "rep 완료 시 요약을 생성하도록 하였다. 종료 시에는 세션 요약이 터미널에 출력된다.",
    size=10, space_after=4)

code_main_loop = '''# 메인 루프 핵심 부분
feedback.update(issues, stage)

if count > prev_count:
    feedback.on_rep_complete(count)
    prev_count = count

draw_status_panel(frame, count, stage, is_normal, issues)
draw_feedback_panel(frame, feedback.get_active_messages(),
                    feedback.get_rep_summary())
draw_session_stats(frame, feedback.get_session_summary())

# 종료 시 세션 요약 출력
print(f"  Total Reps : {counter.count}")
print(f"  Stats      : {feedback.get_session_summary()}")'''
add_code_block(doc, code_main_loop)

# ── 3. 문제점 및 해결 ──
add_styled_paragraph(doc, "3. 문제점 및 해결", size=14, bold=True, space_before=10, space_after=6)

add_styled_paragraph(doc,
    "OpenCV의 cv2.putText 함수는 한글 폰트를 지원하지 않아, 피드백 메시지를 한글로 "
    "표시할 경우 글자가 깨지는 문제가 발생하였다. 이에 따라 모든 피드백 메시지를 영어로 작성하였으며, "
    "향후 PIL(Pillow) 라이브러리의 ImageDraw를 활용하여 한글 렌더링을 지원할 수 있도록 개선할 예정이다.",
    size=10, space_after=4)

add_styled_paragraph(doc,
    "또한 4주차의 judge_squat_pose 함수가 단순 문자열 리스트를 반환하던 것을 "
    "딕셔너리(key, label) 형태로 변경하여 피드백 메시지 매핑이 가능하도록 구조를 개선하였다.",
    size=10, space_after=4)

# ── 4. 향후 계획 ──
add_styled_paragraph(doc, "4. 향후 계획", size=14, bold=True, space_before=10, space_after=6)

plans = [
    "모바일(안드로이드) 환경으로의 전환 준비 (Kotlin + ML Kit 또는 MediaPipe Android)",
    "다양한 운동(런지, 플랭크, 데드리프트 등)에 대한 자세 판별 및 피드백 확장",
    "피드백 메시지 한글화 (PIL/Pillow 연동)",
    "음성 피드백 기능 추가 검토 (TTS 라이브러리 활용)",
    "운동 기록 저장 및 통계 시각화 기능 개발",
]
for plan in plans:
    p = doc.add_paragraph(plan, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ══════════════════════════════════════════════════════════════
# 저장
# ══════════════════════════════════════════════════════════════
save_dir = r"c:\Users\junwoo\Desktop\4학년1학기\종합설계프로젝트\연구노트\제출용"
filename = "연구노트_2026_03_30.docx"
save_path = os.path.join(save_dir, filename)

doc.save(save_path)
print(f"연구노트 생성 완료: {save_path}")
